import os
import streamlit as st
from streamlit_lottie import st_lottie
import requests
from docx import Document
from docx.shared import Pt
import io
from langchain_groq import ChatGroq
from langchain_core.prompts import PromptTemplate
from langchain_core.exceptions import OutputParserException

# Modified environment setup to work with both local and deployment
def get_api_key():
    # First try to get from Streamlit Secrets
    try:
        api_key = st.secrets["GROQ_API_KEY"]
    except:
        # If not found in Streamlit Secrets, try local .env file
        api_key = 'Not Found'
    
    if api_key is None:
        st.error("No API key found. Please set up your GROQ API key in the environment variables.")
        st.stop()
    
    return api_key

class Chain:
    def __init__(self):
        api_key = get_api_key()
        self.llm = ChatGroq(
            temperature=0, 
            groq_api_key=api_key, 
            model_name="llama-3.1-70b-versatile"
        )

    def generate_learning_path(self, educational_background, skills, goals):
        prompt_path = PromptTemplate.from_template(
            """
            ### EDUCATIONAL BACKGROUND:
            {educational_background}

            ### SKILLS:
            {skills}

            ### GOALS:
            {goals}

            ### INSTRUCTION:
            Based on the user's educational background, skills, and specific goals, generate a personalized learning path. For each learning topic, provide the following:

            1. The topic name
            2. Key concepts that will be covered in this topic (5-7 main concepts)
            3. A list of recommended resources categorized as:
                - Books: List 2-3 highly recommended books for the topic (title and author only)
                - Courses: List 2-3 recommended courses or tutorials (include official URLs only for well-known platforms like Coursera, edX, Udemy, LinkedIn Learning)
                - Resources: List 2-3 recommended websites or learning platforms (include official URLs only for well-known platforms)
            4. An estimated time for learning the topic

            IMPORTANT GUIDELINES:
            - Only include URLs for official, well-known educational platforms and resources
            - If you're not completely confident about a URL, provide just the resource name without the link
            - For books, only provide title and author (no links)
            - Each topic must include 5-7 key concepts that will be covered
            
            Structure your response in the following way:
            - **Topic**: [name of the topic]
            - **Estimated Time**: [time estimate]
            - **Key Concepts**:
              - [Concept 1]
              - [Concept 2]
              - [Concept 3]
              - [Concept 4]
              - [Concept 5]
            - **Books**:
              - [Book Title] by [Author]
              - [Book Title] by [Author]
            - **Courses**:
              - [Course Title] - [Platform] ([URL if confident])
              - [Course Title] - [Platform]
            - **Resources**:
              - [Resource Name] ([URL if confident])
              - [Resource Name]

            Remember to maintain high confidence when including URLs and focus on major educational platforms only.
            """
        )
        chain_path = prompt_path | self.llm
        res = chain_path.invoke(input={"educational_background": educational_background, "skills": skills, "goals": goals})
        
        try:
            response = res.content
        except OutputParserException:
            raise OutputParserException("Context too large or output could not be parsed.")
        
        return response

def load_lottieurl(url):
    r = requests.get(url)
    if r.status_code != 200:
        return None
    return r.json()

def generate_docx(learning_path):
    doc = Document()
    doc.add_heading('Personalized Learning Path', 0)
    doc.styles['Normal'].font.size = Pt(12)

    for line in learning_path.split('\n'):
        if line.startswith("- **Topic**"):
            topic = line.split(": ")[1]
            doc.add_paragraph(topic, style="Heading 2")
        elif line.startswith("- **Estimated Time**"):
            time = line.split(": ")[1]
            doc.add_paragraph(f"Estimated Time: {time}")
        elif line.startswith("- **Key Concepts**"):
            doc.add_paragraph("Key Concepts:", style="Heading 3")
        elif line.startswith("- **Books**"):
            doc.add_paragraph("Books:", style="Heading 3")
        elif line.startswith("- **Courses**"):
            doc.add_paragraph("Courses:", style="Heading 3")
        elif line.startswith("- **Resources**"):
            doc.add_paragraph("Resources:", style="Heading 3")
        elif line.startswith("  - "):
            text = line[4:]
            # Create a paragraph with the link if it exists
            if "(" in text and ")" in text and "http" in text:
                name = text[:text.find("(")].strip()
                url = text[text.find("(")+1:text.find(")")].strip()
                p = doc.add_paragraph()
                p.add_run(name + " - ")
                p.add_run(url)
            else:
                doc.add_paragraph(text)
        else:
            doc.add_paragraph(line)
            
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

# Lottie animation URL
lottie_url = "https://assets2.lottiefiles.com/packages/lf20_DMgKk1.json"

# Page config and title
st.set_page_config(
    page_title="Learning Path Generator",
    page_icon="📘",
    layout="wide"
)

# [Previous styling code remains the same]
st.markdown(
    """
    <style>
        /* Base app styling */
        .stApp {
            background-color: #f7ede2;
        }
        
        /* Text input styling */
        .stTextArea textarea {
            background-color: #ffffff;
            color: #000000 !important;
            border: 2px solid #e9c46a;
            border-radius: 8px;
            font-size: 1rem;
            padding: 10px;
        }
        
        /* Labels and text styling */
        .css-1d391kg, .css-10trblm, .css-145kmo2, .css-1629p8f, .css-2trqyj, .css-4yfn50 {
            color: #000000 !important;
            font-weight: 500;
        }
        
        /* Heading styling */
        h1 {
            color: #000000 !important;
            font-weight: 700 !important;
        }
        
        h2, h3 {
            color: #000000 !important;
            font-weight: 600 !important;
        }
        
        /* Button styling */
        .stButton>button {
            background-color: #e9c46a;
            color: #000000;
            border: none;
            border-radius: 8px;
            padding: 0.5em 1em;
            font-weight: bold;
            transition: all 0.3s ease;
        }
        
        .stButton>button:hover {
            background-color: #f4a261;
            color: #000000;
            box-shadow: 0 2px 4px rgba(0,0,0,0.2);
        }
        
        /* Footer styling */
        .footer {
            font-size: 0.9em;
            color: #000000;
            text-align: center;
            margin-top: 20px;
            font-weight: 500;
        }
        
        /* Markdown text */
        .stMarkdown {
            color: #000000 !important;
        }
        
        /* Warning and error messages */
        .stAlert {
            color: #000000 !important;
        }
        
        /* Spinner text */
        .stSpinner {
            color: #000000 !important;
        }
        
        /* Text area labels */
        .stTextArea label {
            color: #000000 !important;
            font-weight: 600 !important;
        }
        
        /* Placeholder text */
        .stTextArea textarea::placeholder {
            color: #666666 !important;
        }
        
        /* Download button */
        .stDownloadButton>button {
            background-color: #e9c46a;
            color: #000000;
            font-weight: bold;
        }
        
        .stDownloadButton>button:hover {
            background-color: #f4a261;
            color: #000000;
        }
        
        /* Lottie container styling */
        .lottie-container {
            background: transparent !important;
            padding: 20px;
            margin: 20px auto;
            max-width: 800px;
            border-radius: 12px;
            display: flex;
            justify-content: center;
            align-items: center;
        }
        
        /* Override stLottie default styling */
        .stLottie {
            background: transparent !important;
        }
        
        .stLottie > div {
            background: transparent !important;
        }

        /* Sidebar styling */
        .css-1d391kg {
            background-color: #f7ede2 !important;
        }
        
        [data-testid="stSidebar"] {
            background-color: #f7ede2;
            border-right: 2px solid #e9c46a;
        }
        
        [data-testid="stSidebar"] > div:first-child {
            background-color: #f7ede2;
        }
        
        /* Sidebar text and headings */
        [data-testid="stSidebar"] .css-10trblm {
            color: #000000 !important;
            font-weight: 700 !important;
        }
        
        [data-testid="stSidebar"] p {
            color: #000000 !important;
            font-weight: 500 !important;
        }
        
        [data-testid="stSidebar"] h1,
        [data-testid="stSidebar"] h2,
        [data-testid="stSidebar"] h3 {
            color: #000000 !important;
            font-weight: 600 !important;
        }
        
        [data-testid="stSidebar"] ul {
            color: #000000 !important;
        }
        
        [data-testid="stSidebar"] li {
            color: #000000 !important;
        }
        
        /* Error message styling */
        .stException {
            color: #000000 !important;
            background-color: #ffe8e8 !important;
            padding: 1rem !important;
            border-radius: 8px !important;
        }
    </style>
    """, unsafe_allow_html=True
)

# Sidebar with instructions
with st.sidebar:
    st.markdown("""
    # 📚 How to Use
    
    Welcome to the Learning Path Generator! Follow these steps to get your personalized learning path:
    
    ### 1. Educational Background 📘
    * Enter your current education level
    * Include your major/field of study
    * List relevant courses you've taken
    * Mention any certifications
    
    ### 2. Skills 🔧
    * List your technical skills
    * Include programming languages
    * Add tools you're familiar with
    * Mention soft skills
    
    ### 3. Goals 🎯
    * Specify your career objectives
    * Mention target roles/positions
    * Include time frame for goals
    * List specific areas you want to learn
    
    ### 4. Generate & Download 📥
    * Click "Generate Learning Path"
    * Review the customized path
    * Download as a Word document
    * Follow the suggested resources
    
    ### Tips for Best Results ✨
    * Be specific in your descriptions
    * Include skill levels where relevant
    * Clearly state your timeline
    * Mention any preferences for learning style
    
    Need help? Contact: shguda@syr.edu (or) bdasari@syr.edu
    """)

# Main content
st.title("📘 LEAP - Learning Enhancement And Progression: Personalized Learning Path Generator")
st.write("Enter your educational background, skills, and goals to generate a customized learning path.")

# Initialize Chain
chain = Chain()

# Lottie animation display
lottie_animation = load_lottieurl(lottie_url)
if lottie_animation:
    st.markdown('<div class="lottie-container">', unsafe_allow_html=True)
    st_lottie(
        lottie_animation,
        height=200,
        key="motivation",
        quality="high",
        speed=1
    )
    st.markdown('</div>', unsafe_allow_html=True)

# User inputs
educational_background = st.text_area("📘 Educational Background", "", height=100, placeholder="Enter details (e.g., degree, major, relevant courses):")
skills = st.text_area("🔧 Skills", "", height=100, placeholder="Enter details (e.g., programming languages, technical skills, tools):")
goals = st.text_area("🎯 Goals", "", height=100, placeholder="Enter details (e.g., career objectives, learning goals):")

# Generate learning path button
if st.button("Generate Learning Path"):
    if educational_background and skills and goals:
        with st.spinner("Generating your personalized learning path..."):
            try:
                learning_path = chain.generate_learning_path(educational_background, skills, goals)
                st.session_state.learning_path = learning_path
                st.subheader("Your Customized Learning Path:")
                st.markdown("<hr>", unsafe_allow_html=True)
                for module in learning_path.split('\n\n'):
                    st.write(module)
                    st.markdown("<hr>", unsafe_allow_html=True)
                docx_file = generate_docx(learning_path)
                st.download_button(
                    label="Download Learning Path as .docx",
                    data=docx_file.getvalue(),
                    file_name="learning_path.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception as e:
                st.error(f"An error occurred: {e}")
    else:
        st.warning("Please provide educational background, skills, and goals.")

# Footer
st.markdown("<div class='footer'>Made with ❤️ by Shashank Guda & Bhumika Dasari</div>", unsafe_allow_html=True)